from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
h1 = doc.styles['Heading 1']
h1.font.size = Pt(18)
h1.font.color.rgb = RGBColor(0x8F, 0x1D, 0x31)
h2 = doc.styles['Heading 2']
h2.font.size = Pt(14)
h2.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

# Title Page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Eden Relics')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x8F, 0x1D, 0x31)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Security Assessment Summary')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('INITIAL ASSESSMENT — PRE-REMEDIATION\n').bold = True
meta.add_run(f'Date: {datetime.now().strftime("%d %B %Y")}\n')
meta.add_run('Classification: CONFIDENTIAL\n')
meta.add_run('Prepared by: Security Assessment Team')

doc.add_page_break()

# Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'A comprehensive security assessment was conducted on the Eden Relics e-commerce platform, '
    'covering the backend application, payment processing workflows, banking integrations, and '
    'personal data handling practices. The assessment combined static analysis with dynamic testing '
    'to identify potential security risks.'
)
doc.add_paragraph(
    'The platform demonstrates a solid security foundation with industry-standard authentication, '
    'encrypted communications, security headers, and rate limiting. A number of areas for improvement '
    'were identified, ranging from items requiring immediate attention to longer-term best practice '
    'enhancements. A remediation plan has been agreed and work is underway.'
)

doc.add_heading('Assessment Overview', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = 'Priority'
headers[1].text = 'Findings'
headers[2].text = 'Status'
data = [
    ('Immediate', '6', 'Remediation in progress'),
    ('High', '8', 'Scheduled for remediation'),
    ('Medium', '12', 'Planned for upcoming sprint'),
    ('Low', '5', 'Tracked for future improvement'),
    ('Informational', '4', 'Best practice recommendations'),
]
for i, (sev, count, status) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = sev
    row[1].text = count
    row[2].text = status

doc.add_paragraph()

# Scope
doc.add_heading('2. Scope & Methodology', level=1)
doc.add_paragraph(
    'The assessment covered the full backend application including all API endpoints, service '
    'integrations, data storage, and application configuration. Testing methodology combined '
    'code review with endpoint-level testing.'
)
doc.add_heading('Focus Areas', level=2)
doc.add_paragraph('Banking integration security (OAuth flows, credential management, data sync)', style='List Bullet')
doc.add_paragraph('Payment processing integrity (checkout flow, webhook handling, order security)', style='List Bullet')
doc.add_paragraph('Personal data protection (storage practices, access controls, GDPR alignment)', style='List Bullet')
doc.add_paragraph('Authentication and session management', style='List Bullet')
doc.add_paragraph('Input validation and injection prevention', style='List Bullet')
doc.add_paragraph('Infrastructure and configuration security', style='List Bullet')

doc.add_page_break()

# Findings by Category
doc.add_heading('3. Findings — Immediate Priority', level=1)

doc.add_heading('3.1 Credential Management', level=2)
doc.add_paragraph(
    'Development configuration was found to contain service credentials that should be managed '
    'through a dedicated secrets management solution. A process has been initiated to migrate '
    'credentials to environment-based configuration and rotate affected values.'
)

doc.add_heading('3.2 Access Control on Customer Data', level=2)
doc.add_paragraph(
    'An access control gap was identified in a customer-facing data retrieval endpoint. The endpoint '
    'did not fully enforce ownership verification, meaning it was theoretically possible to access '
    'records belonging to other users. This has been prioritised for immediate remediation.'
)

doc.add_heading('3.3 Token Generation Strength', level=2)
doc.add_paragraph(
    'Several security-sensitive tokens (used for account verification and password recovery) were '
    'generated using a method that does not meet current cryptographic best practice standards. '
    'These are being replaced with tokens generated using a cryptographically secure random number '
    'generator.'
)

doc.add_heading('3.4 Third-Party Account Linking', level=2)
doc.add_paragraph(
    'The social login integration was found to automatically link third-party accounts to existing '
    'accounts based on email address without requiring additional verification. Additional safeguards '
    'are being implemented to ensure account linking requires proper authorisation.'
)

doc.add_heading('3.5 Sensitive Credential Storage', level=2)
doc.add_paragraph(
    'Credentials used for banking integration were found to be stored without application-level '
    'encryption. While database-level security controls exist, additional encryption at rest is '
    'being implemented as a defence-in-depth measure.'
)

doc.add_heading('3.6 Error Handling', level=2)
doc.add_paragraph(
    'One integration endpoint was found to return detailed internal error information in its '
    'responses. This has been updated to return generic error messages while retaining detailed '
    'logging for internal diagnostics.'
)

doc.add_page_break()

doc.add_heading('4. Findings — High Priority', level=1)

doc.add_heading('4.1 Input Validation', level=2)
doc.add_paragraph(
    'Several data input points were found to lack comprehensive validation rules, including '
    'minimum complexity requirements for passwords and format checks for email addresses. '
    'Validation rules are being added across all relevant input models.'
)

doc.add_heading('4.2 Rate Limiting Coverage', level=2)
doc.add_paragraph(
    'Certain sensitive endpoints (such as password recovery) were not covered by the existing '
    'rate limiting infrastructure. Rate limiting is being extended to cover all sensitive operations.'
)

doc.add_heading('4.3 Payment Webhook Handling', level=2)
doc.add_paragraph(
    'The payment webhook handler was found to lack idempotency protection. In the unlikely event '
    'of duplicate webhook delivery, processing could be repeated. An idempotency mechanism is being '
    'implemented.'
)

doc.add_heading('4.4 Banking Integration OAuth Flow', level=2)
doc.add_paragraph(
    'The OAuth integration with the banking provider was missing a cross-site request forgery '
    'protection check during the callback phase. This is being addressed by validating the state '
    'parameter throughout the flow.'
)

doc.add_heading('4.5 Data Encryption at Rest', level=2)
doc.add_paragraph(
    'Personal data including customer addresses and account security credentials are stored without '
    'application-level encryption. While the database is access-controlled, additional encryption '
    'is recommended as a defence-in-depth measure and for GDPR alignment.'
)

doc.add_heading('4.6 Session Lifetime', level=2)
doc.add_paragraph(
    'Authentication session durations were found to be longer than industry-recommended values. '
    'A plan is in place to reduce session lifetimes and implement a more robust token refresh mechanism.'
)

doc.add_heading('4.7 Email Template Security', level=2)
doc.add_paragraph(
    'User-provided content in certain email templates was not being properly encoded before '
    'rendering, creating a theoretical content injection risk. Output encoding has been applied.'
)

doc.add_heading('4.8 Development Diagnostics', level=2)
doc.add_paragraph(
    'A diagnostic endpoint used during development was found to expose operational metadata. '
    'This endpoint is protected by administrative access controls but is flagged for removal '
    'prior to production deployment.'
)

doc.add_page_break()

doc.add_heading('5. Findings — Medium & Lower Priority', level=1)
doc.add_paragraph(
    'The following areas were identified as opportunities for further hardening. None represent '
    'immediately exploitable vulnerabilities but address defence-in-depth and compliance best practices.'
)

medium_items = [
    'Account enumeration resistance on registration and recovery flows',
    'Cross-origin resource sharing policy refinement',
    'Proxy header trust configuration for rate limiting',
    'Audit logging for financial data operations',
    'Enhanced file upload content validation',
    'Input range validation on financial data entry',
    'Multi-factor authentication attempt limiting',
    'Email format validation on guest checkout',
    'Subscriber email validation improvements',
    'Token refresh window reduction',
    'Browser permissions policy expansion',
    'Third-party API error log sanitisation',
]

for item in medium_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Additional informational items include session revocation capabilities, '
    'data type validation on status fields, deployment pipeline security, API versioning strategy, '
    'and audit logging for GDPR Article 30 compliance.')

doc.add_page_break()

# Area Deep Dives
doc.add_heading('6. Banking Integration Summary', level=1)
doc.add_paragraph(
    'The banking integration uses industry-standard OAuth 2.0 for authorisation and communicates '
    'exclusively over encrypted channels. Transaction data is synchronised automatically with '
    'deduplication controls in place. Areas for improvement include strengthening the OAuth CSRF '
    'protection, adding encryption for stored credentials, and removing development diagnostic '
    'functionality before production release.'
)

doc.add_heading('7. Payment Processing Summary', level=1)
doc.add_paragraph(
    'The payment flow correctly sources pricing from authoritative server-side records rather '
    'than client input, and webhook signatures are properly validated. Areas for improvement '
    'include adding idempotency protection to webhook processing and strengthening access controls '
    'on order retrieval.'
)

doc.add_heading('8. Personal Data Summary', level=1)
doc.add_paragraph(
    'The platform stores personal data necessary for order fulfilment and account management. '
    'Database access is controlled and communications are encrypted in transit. Areas for '
    'improvement include implementing application-level encryption at rest for sensitive fields, '
    'adding audit logging for administrative data access, and strengthening access control '
    'granularity on customer-facing endpoints.'
)

doc.add_page_break()

# Remediation
doc.add_heading('9. Remediation Roadmap', level=1)

doc.add_heading('Immediate Actions (In Progress)', level=2)
doc.add_paragraph('Credential rotation and migration to secure configuration management', style='List Bullet')
doc.add_paragraph('Access control enforcement on customer data endpoints', style='List Bullet')
doc.add_paragraph('Cryptographic token generation upgrade', style='List Bullet')
doc.add_paragraph('Third-party account linking safeguards', style='List Bullet')
doc.add_paragraph('Error response sanitisation', style='List Bullet')
doc.add_paragraph('Email template output encoding', style='List Bullet')

doc.add_heading('Short Term', level=2)
doc.add_paragraph('Input validation hardening across all data entry points', style='List Bullet')
doc.add_paragraph('Rate limiting extension to all sensitive operations', style='List Bullet')
doc.add_paragraph('Payment webhook idempotency implementation', style='List Bullet')
doc.add_paragraph('Banking OAuth CSRF protection', style='List Bullet')
doc.add_paragraph('Cross-origin policy tightening', style='List Bullet')

doc.add_heading('Medium Term', level=2)
doc.add_paragraph('Application-level encryption at rest for personal data', style='List Bullet')
doc.add_paragraph('Banking credential encryption', style='List Bullet')
doc.add_paragraph('Session lifetime reduction and refresh token improvements', style='List Bullet')
doc.add_paragraph('Comprehensive audit logging', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('END OF INITIAL ASSESSMENT').bold = True

doc.save(r'C:\Users\peter\source\eden-relics-mono\Eden_Relics_Security_Assessment_Initial_Redacted.docx')
print('Redacted initial report generated.')
