from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

title_style = doc.styles['Title']
title_style.font.size = Pt(28)
title_style.font.color.rgb = RGBColor(0x8F, 0x1D, 0x31)

h1 = doc.styles['Heading 1']
h1.font.size = Pt(18)
h1.font.color.rgb = RGBColor(0x8F, 0x1D, 0x31)

h2 = doc.styles['Heading 2']
h2.font.size = Pt(14)
h2.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

h3 = doc.styles['Heading 3']
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(0x5A, 0x58, 0x58)

# Title Page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Eden Relics')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x8F, 0x1D, 0x31)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Security Audit & Penetration Test Report')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('INITIAL FINDINGS — PRE-REMEDIATION\n').bold = True
meta.add_run(f'Date: {datetime.now().strftime("%d %B %Y")}\n')
meta.add_run('Classification: CONFIDENTIAL\n')
meta.add_run('Prepared by: Security Audit Team')

doc.add_page_break()

# Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'This report presents the findings of a comprehensive security audit and penetration test '
    'conducted on the Eden Relics e-commerce backend application (ASP.NET / C# / .NET 10). '
    'The audit focused on three critical areas: banking integration (Monzo API), product purchasing '
    '(Stripe payments), and personal data handling (GDPR compliance).'
)
doc.add_paragraph(
    'The application demonstrates solid foundational security practices including JWT authentication, '
    'HTTPS enforcement, security headers, rate limiting infrastructure, and proper password hashing '
    '(PBKDF2-SHA256). However, several critical and high-severity vulnerabilities were identified '
    'that require immediate remediation before the application should handle live financial transactions '
    'or personal data.'
)

# Stats
doc.add_heading('Vulnerability Summary', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = 'Severity'
headers[1].text = 'Count'
headers[2].text = 'Status'
data = [
    ('CRITICAL', '6', 'Requires immediate remediation'),
    ('HIGH', '8', 'Requires remediation before production'),
    ('MEDIUM', '12', 'Should be addressed in next sprint'),
    ('LOW', '5', 'Address when convenient'),
    ('INFORMATIONAL', '4', 'Best practice recommendations'),
]
for i, (sev, count, status) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = sev
    row[1].text = count
    row[2].text = status

doc.add_paragraph()

# Scope
doc.add_heading('2. Scope & Methodology', level=1)
doc.add_paragraph(
    'The audit covered the complete ASP.NET backend codebase including 17 API controllers, '
    '8 service classes, data access layer, and application configuration. Testing methodology '
    'combined static code analysis (SAST) with dynamic endpoint testing against the integration '
    'test infrastructure.'
)
doc.add_heading('Focus Areas', level=2)
bullets = doc.add_paragraph()
bullets.style = 'List Bullet'
bullets.text = 'Banking Integration — Monzo OAuth flow, token storage, transaction sync, API security'
doc.add_paragraph('Product Purchasing — Stripe checkout, webhook handling, price integrity, order access control', style='List Bullet')
doc.add_paragraph('Personal Data — PII storage, encryption at rest, GDPR compliance, data exposure', style='List Bullet')
doc.add_paragraph('Authentication & Authorization — JWT tokens, password handling, MFA, OAuth providers', style='List Bullet')
doc.add_paragraph('Input Validation — Injection attacks, XSS, IDOR, mass assignment', style='List Bullet')
doc.add_paragraph('Infrastructure — CORS, security headers, rate limiting, secrets management', style='List Bullet')

doc.add_page_break()

# Critical Findings
doc.add_heading('3. Critical Findings', level=1)

# C1
doc.add_heading('C1: Hardcoded API Secrets in Source Control', level=2)
doc.add_paragraph('File: appsettings.Development.json | Severity: CRITICAL')
doc.add_paragraph(
    'Production API credentials are committed to the repository in plaintext, including the Monzo '
    'OAuth client secret, Anthropic API key, Google OAuth client ID, and database credentials. '
    'Anyone with repository access can obtain these credentials and access financial APIs, '
    'AI services, and the database directly.'
)
doc.add_paragraph('Impact: ', style='List Bullet').add_run('Unauthorized access to Monzo banking API — potential fraudulent transactions').bold = False
doc.add_paragraph('Impact: ', style='List Bullet').add_run('Anthropic API key abuse — cost exposure').bold = False
doc.add_paragraph('Impact: ', style='List Bullet').add_run('Database credential compromise').bold = False
doc.add_paragraph('Recommendation: Move all secrets to environment variables or a secrets manager (Azure Key Vault, dotnet user-secrets). Immediately rotate all exposed credentials.')

# C2
doc.add_heading('C2: Unauthenticated Order Access (IDOR)', level=2)
doc.add_paragraph('File: OrdersController.cs:223-235 | Severity: CRITICAL')
doc.add_paragraph(
    'The GET /api/orders/{id} endpoint has no [Authorize] attribute and no ownership check. '
    'Any unauthenticated user who knows or guesses an order GUID can retrieve the full order '
    'details including customer email, shipping/billing addresses, purchased items, and prices. '
    'While GUIDs are not sequential, they are returned in API responses and may be cached or logged.'
)
doc.add_paragraph('Recommendation: Add [Authorize] attribute and verify the requesting user owns the order or is an Admin.')

# C3
doc.add_heading('C3: Weak Security Tokens (GUID-based)', level=2)
doc.add_paragraph('File: AuthController.cs:51-52, 137-138, 209-210 | Severity: CRITICAL')
doc.add_paragraph(
    'Email verification tokens, password reset tokens, and other security tokens are generated '
    'from System.Guid, which is not cryptographically secure. GUIDs are based on predictable '
    'algorithms (timestamp + MAC address or random, depending on version) and produce only 16 bytes '
    'of entropy. An attacker who knows the approximate creation time could narrow the search space.'
)
doc.add_paragraph('Recommendation: Use System.Security.Cryptography.RandomNumberGenerator to generate 32-byte tokens.')

# C4
doc.add_heading('C4: OAuth Account Linking Without Verification', level=2)
doc.add_paragraph('File: AuthController.cs:242-253 | Severity: CRITICAL')
doc.add_paragraph(
    'When a user logs in via an external OAuth provider (Google, Facebook, Apple), if no account '
    'exists with that provider ID but an account exists with the same email, the external account '
    'is automatically linked without any user confirmation. An attacker could create a Google account '
    'with the victim\'s email and gain full access to their Eden Relics account.'
)
doc.add_paragraph('Recommendation: Require explicit user authorization to link external accounts, or require the existing account\'s password before linking.')

# C5
doc.add_heading('C5: Monzo Tokens Stored in Plaintext', level=2)
doc.add_paragraph('File: MonzoService.cs, MonzoToken entity | Severity: CRITICAL')
doc.add_paragraph(
    'Monzo OAuth access tokens and refresh tokens are stored in the database as plaintext strings. '
    'A database compromise would grant the attacker direct access to the connected Monzo business '
    'account, including the ability to read all transactions, account balances, and potentially '
    'initiate actions via the Monzo API.'
)
doc.add_paragraph('Recommendation: Encrypt tokens at rest using AES-256-GCM with keys stored in a separate key management service.')

# C6
doc.add_heading('C6: Exception Details Exposed to Clients', level=2)
doc.add_paragraph('File: MonzoController.cs:198 | Severity: CRITICAL')
doc.add_paragraph(
    'The Monzo sync endpoint returns raw exception messages and inner exception details to the client: '
    'return StatusCode(500, new { error = ex.Message, inner = ex.InnerException?.Message }). '
    'This can leak internal implementation details, database connection strings, file paths, and '
    'stack traces to attackers.'
)
doc.add_paragraph('Recommendation: Log exceptions server-side and return generic error messages to clients.')

doc.add_page_break()

# High Findings
doc.add_heading('4. High Severity Findings', level=1)

doc.add_heading('H1: Missing Input Validation on Authentication DTOs', level=2)
doc.add_paragraph('Files: AuthDtos.cs, AccountDtos.cs, OrderDtos.cs | Severity: HIGH')
doc.add_paragraph(
    'Critical DTOs including RegisterDto, ChangePasswordDto, and AddressDto lack validation attributes. '
    'Passwords have no minimum length or complexity requirements. Email fields lack format validation. '
    'Name fields have no length limits, allowing extremely long strings that could cause storage or '
    'processing issues.'
)

doc.add_heading('H2: Forgot-Password Endpoint Not Rate Limited', level=2)
doc.add_paragraph('File: AuthController.cs:126 | Severity: HIGH')
doc.add_paragraph(
    'The forgot-password endpoint lacks rate limiting, allowing an attacker to flood a victim\'s '
    'inbox with password reset emails or exhaust the email service quota.'
)

doc.add_heading('H3: Stripe Webhook Lacks Idempotency Check', level=2)
doc.add_paragraph('File: OrdersController.cs:155-206 | Severity: HIGH')
doc.add_paragraph(
    'The Stripe webhook handler processes checkout.session.completed events without checking if '
    'the event has already been processed. If Stripe retries the webhook (e.g., due to a timeout), '
    'the order could be marked as Paid twice, inventory double-sold, or other side effects triggered.'
)

doc.add_heading('H4: Monzo OAuth State Parameter Not Validated', level=2)
doc.add_paragraph('File: MonzoController.cs:36-77 | Severity: HIGH')
doc.add_paragraph(
    'The OAuth state parameter is generated in the Connect endpoint but never validated in the '
    'Callback endpoint. This allows CSRF attacks where an attacker could substitute their own '
    'authorization code and link their Monzo account to the admin\'s session.'
)

doc.add_heading('H5: No PII Encryption at Rest', level=2)
doc.add_paragraph('Files: User.cs, Order.cs | Severity: HIGH')
doc.add_paragraph(
    'Personal data including delivery addresses, billing addresses, guest emails, and MFA secrets '
    'are stored in plaintext in the database. Under GDPR and PCI-DSS, sensitive personal data '
    'should be encrypted at rest.'
)

doc.add_heading('H6: JWT Token Lifetime Excessive (7 Days)', level=2)
doc.add_paragraph('File: AuthController.cs:387 | Severity: HIGH')
doc.add_paragraph(
    'JWT tokens are valid for 7 days with a 30-day refresh window. If a token is compromised, '
    'an attacker has extended access. Industry standard is 15-60 minutes with short-lived refresh tokens.'
)

doc.add_heading('H7: HTML Injection in Contact Email', level=2)
doc.add_paragraph('File: EmailService.cs:104-114 | Severity: HIGH')
doc.add_paragraph(
    'User-supplied content (name, subject, message) from the contact form is interpolated directly '
    'into HTML email templates without encoding. An attacker could inject arbitrary HTML including '
    'phishing links or malicious content that renders in the admin\'s email client.'
)

doc.add_heading('H8: Debug Endpoint Exposes Banking Metadata', level=2)
doc.add_paragraph('File: MonzoController.cs:116-142 | Severity: HIGH')
doc.add_paragraph(
    'The /api/monzo/debug-transactions endpoint returns sensitive information including account IDs '
    'and token expiry times. While protected by Admin auth, this endpoint should not exist in production.'
)

doc.add_page_break()

# Medium Findings
doc.add_heading('5. Medium Severity Findings', level=1)

medium_findings = [
    ('M1: Email Enumeration via Registration', 'AuthController.cs:45-48',
     'Registration returns 409 Conflict for existing emails, allowing attackers to enumerate valid accounts.'),
    ('M2: CORS AllowAnyMethod/AllowAnyHeader', 'Program.cs:140-142',
     'CORS policy allows all HTTP methods and headers. Should restrict to specific methods (GET, POST, PUT, DELETE, PATCH, OPTIONS).'),
    ('M3: IP Spoofing in Rate Limiting', 'Program.cs:33-35',
     'Rate limiting trusts X-Forwarded-For header directly. Attackers can bypass rate limits by spoofing headers.'),
    ('M4: No Audit Logging for Financial Operations', 'FinanceController.cs, MonzoController.cs',
     'Financial transaction modifications and data exports have no audit trail, making it impossible to detect insider threats.'),
    ('M5: Receipt Upload MIME Type Not Validated', 'FinanceController.cs:101-121',
     'File upload validates extension only, not actual MIME type. Attacker could upload malicious files with renamed extensions.'),
    ('M6: Transaction Amount Not Range-Validated', 'FinanceController.cs:43-61',
     'Financial transaction amounts have no range validation. Arbitrarily large values could manipulate reports.'),
    ('M7: MFA Code Verification Not Rate Limited', 'AccountController.cs:158-184',
     'TOTP verification has no attempt counter or rate limiting, allowing brute force of 6-digit codes.'),
    ('M8: Guest Email Format Not Validated', 'OrdersController.cs:33-35',
     'Guest checkout accepts any non-empty string as email. Only whitespace is checked, not email format.'),
    ('M9: Mailing List Email Validation Weak', 'MailingListController.cs:16',
     'Email validation only checks Contains("@") and Length >= 5. Allows many invalid email formats.'),
    ('M10: Token Refresh Window Too Long', 'AuthController.cs:480',
     'Expired tokens can be refreshed up to 30 days after expiry, defeating the purpose of token expiration.'),
    ('M11: Permissions-Policy Incomplete', 'Program.cs:177',
     'Only restricts camera, microphone, and geolocation. Missing payment, usb, and other permission restrictions.'),
    ('M12: Monzo Error Responses Logged Verbatim', 'MonzoService.cs:43-44',
     'Full API error response bodies are logged, potentially including sensitive information.'),
]

table = doc.add_table(rows=len(medium_findings) + 1, cols=3)
table.style = 'Light Grid Accent 1'
headers = table.rows[0].cells
headers[0].text = 'ID'
headers[1].text = 'Location'
headers[2].text = 'Description'
for i, (fid, loc, desc) in enumerate(medium_findings):
    row = table.rows[i + 1].cells
    row[0].text = fid
    row[1].text = loc
    row[2].text = desc

doc.add_page_break()

# Low / Info
doc.add_heading('6. Low Severity & Informational Findings', level=1)

low_findings = [
    ('L1', 'No token revocation mechanism — compromised tokens cannot be invalidated before expiry'),
    ('L2', 'Order status accepts any string — no enum validation (Pending, Paid, Shipped, etc.)'),
    ('L3', 'Auto-migrations run on startup in production — risky for availability'),
    ('L4', 'Admin promote endpoint exists (dev-only) — pattern risk if environment misconfigured'),
    ('L5', 'No API versioning strategy — breaking changes affect all clients'),
    ('I1', 'Password hashing uses PBKDF2-SHA256 (acceptable) — could upgrade to Argon2id'),
    ('I2', 'CSP header is very restrictive (default-src none) — correct for API-only backend'),
    ('I3', 'HSTS header applied in all environments including development'),
    ('I4', 'No structured audit logging framework — consider adding for GDPR Article 30 compliance'),
]

for fid, desc in low_findings:
    doc.add_paragraph(f'{fid}: {desc}', style='List Bullet')

doc.add_page_break()

# Banking Integration Deep Dive
doc.add_heading('7. Banking Integration Analysis (Monzo)', level=1)
doc.add_paragraph(
    'The Monzo banking integration connects to the Monzo API via OAuth 2.0 to sync business '
    'transactions automatically. The following specific risks were identified:'
)
doc.add_heading('OAuth Flow', level=2)
doc.add_paragraph('The OAuth implementation generates a state parameter for CSRF protection but does not validate it on callback (C4). The authorization code exchange correctly uses server-side requests with client credentials.')
doc.add_heading('Token Management', level=2)
doc.add_paragraph('Access and refresh tokens are stored in plaintext in the database (C5). Token refresh is handled automatically with a 5-minute buffer before expiry. No encryption, no token rotation logging, and no alerting on refresh failures.')
doc.add_heading('Transaction Sync', level=2)
doc.add_paragraph('The background sync service runs hourly and deduplicates by Monzo transaction ID. However, there is no database unique constraint on MonzoId, creating a race condition risk for duplicates. Transaction data integrity is not validated before storage.')
doc.add_heading('Debug Endpoint', level=2)
doc.add_paragraph('A /debug-transactions endpoint exposes account metadata and token expiry information (H8). This should be removed before production deployment.')

doc.add_page_break()

# Purchasing Deep Dive
doc.add_heading('8. Product Purchasing Analysis (Stripe)', level=1)
doc.add_paragraph(
    'The purchasing flow creates orders, generates Stripe Checkout sessions, and handles '
    'payment completion via webhooks.'
)
doc.add_heading('Price Integrity', level=2)
doc.add_paragraph('Prices are correctly fetched from the database at order creation time (not from client input). The Stripe session uses database-fetched prices. However, there is no idempotency check on webhook processing (H3).')
doc.add_heading('Order Access Control', level=2)
doc.add_paragraph('The most critical finding: GET /api/orders/{id} is completely unauthenticated (C2). Any person can retrieve any order by GUID, exposing customer PII, order details, and potentially shipping addresses.')
doc.add_heading('Webhook Security', level=2)
doc.add_paragraph('Stripe webhook signature validation is correctly implemented using EventUtility.ConstructEvent. However, the handler lacks idempotency checks — if Stripe retries a webhook, the order processing runs again.')

doc.add_page_break()

# Personal Data Deep Dive
doc.add_heading('9. Personal Data Analysis (GDPR)', level=1)
doc.add_paragraph(
    'The application stores personal data including names, emails, addresses, payment card '
    'references, and MFA secrets. The following GDPR-relevant issues were identified:'
)
doc.add_heading('Data Storage', level=2)
doc.add_paragraph('All PII is stored in plaintext in the PostgreSQL database. Under GDPR Article 32, appropriate technical measures (including encryption) must be implemented to protect personal data. Addresses, MFA secrets, and guest emails lack encryption at rest (H5).')
doc.add_heading('Data Access', level=2)
doc.add_paragraph('The AdminUsersController exposes all user data (emails, names, mailing list status, favourites) to any admin user. No field-level access control or audit logging exists.')
doc.add_heading('Data Breach Risk', level=2)
doc.add_paragraph('The combination of the IDOR vulnerability (C2) with plaintext PII storage means a data breach could expose customer orders and personal data without any authentication required.')

doc.add_page_break()

# Recommendations
doc.add_heading('10. Remediation Priority', level=1)

doc.add_heading('Immediate (Before Next Deployment)', level=2)
doc.add_paragraph('1. Rotate all exposed credentials (Monzo, Anthropic, Google OAuth, database)', style='List Bullet')
doc.add_paragraph('2. Add [Authorize] and ownership check to OrdersController.GetById', style='List Bullet')
doc.add_paragraph('3. Replace GUID tokens with cryptographically secure random tokens', style='List Bullet')
doc.add_paragraph('4. Remove or require password confirmation for OAuth account linking', style='List Bullet')
doc.add_paragraph('5. Remove exception details from Monzo sync error responses', style='List Bullet')
doc.add_paragraph('6. HTML-encode user input in contact email templates', style='List Bullet')

doc.add_heading('Short Term (Next Sprint)', level=2)
doc.add_paragraph('7. Add input validation to all DTOs (password complexity, email format, field lengths)', style='List Bullet')
doc.add_paragraph('8. Add rate limiting to forgot-password endpoint', style='List Bullet')
doc.add_paragraph('9. Add webhook idempotency check for Stripe', style='List Bullet')
doc.add_paragraph('10. Validate Monzo OAuth state parameter', style='List Bullet')
doc.add_paragraph('11. Restrict CORS methods', style='List Bullet')
doc.add_paragraph('12. Restrict CORS headers', style='List Bullet')

doc.add_heading('Medium Term', level=2)
doc.add_paragraph('13. Implement PII encryption at rest', style='List Bullet')
doc.add_paragraph('14. Encrypt Monzo tokens at rest', style='List Bullet')
doc.add_paragraph('15. Reduce JWT lifetime and improve refresh token handling', style='List Bullet')
doc.add_paragraph('16. Add audit logging for sensitive operations', style='List Bullet')
doc.add_paragraph('17. Implement token revocation', style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('END OF INITIAL REPORT').bold = True

doc.save(r'C:\Users\peter\source\eden-relics-mono\Eden_Relics_Security_Audit_Initial.docx')
print('Initial report generated successfully.')
